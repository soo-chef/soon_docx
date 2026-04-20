"""
영양사정기록지 docx 자동 채우기 모듈

테이블 행/셀 구조 (unique cells 기준):
  행01 (8셀): 성명 | [값] | 입소일 | [값] | 생년월일 | [값] | 성별 | [값]
  행02 (4셀): 작성일 | [값] | 작성자 | [값]
  행04 (8셀): 신장 | [값 cm] | 평소체중 | [값 kg] | 등급 | [값] | 식사유형 | [값]
  행06 (4셀): 1일필요열량 | [값 kcal] | 1일필요단백질 | [값 g]
  행08 (4셀): 식사방법 | [체크박스] | 식사섭취상태 | [체크박스]
  행09 (4셀): 식사속도 | [체크박스] | 도구사용 | [체크박스 다중]
  행10 (2셀): 식사시문제점 | [체크박스 다중]
  행12 (2셀): 치아상태 | [체크박스]
  행13 (2셀): 소화기능 | [체크박스]
  행14 (2셀): 배설양상 | [체크박스]
  행15 (2셀): 특이체질 | [체크박스+내용]
  행17 (4셀): 선호음식 | [값] | 비선호음식 | [값]
  행18 (2셀): 식품알러지 | [체크박스+내용]
  행20 (2셀): 주요진단명 | [값]
  행21 (2셀): 주요질환 | [체크박스 다중+기타내용]
  행22 (2셀): 현재복용약물 | [체크박스+내용]
  행23 (2셀): 영양관련약물영향 | [체크박스 다중+기타내용]
  행25 (4셀): 종교 | [체크박스] | 금식일기도시간 | [값]
  행26 (2셀): 종교적식사제한 | [체크박스+내용]
  행27 (2셀): 문화적식습관 | [체크박스+내용]
  행28 (2셀): 출신지역특성 | [체크박스+내용]
  행30 (3셀): 개별욕구제목 | 수급자 | [값]
  행31 (3셀): 개별욕구제목 | 보호자 | [값]
  행33 (2셀): 영양사총평 | [값]
  행34 (1셀): 구역 제목(첨부)
  행35 (2셀): 식사사진등첨부(라벨) | [이미지 — `식사사진등첨부`·`식사사진등첨부2` URL, 2열이면 같은 줄에 나란히·크기 자동]
"""
import datetime
import io
import os
import platform
import re
import shutil
import tempfile

from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, '영양사정기록지_개정.docx')

# 로컬(Windows): output/ 폴더 사용 / 클라우드(Linux): 시스템 임시 폴더 사용
if platform.system() == 'Windows':
    OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
    os.makedirs(OUTPUT_DIR, exist_ok=True)
else:
    OUTPUT_DIR = tempfile.gettempdir()

# ─────────────────────────────────────────
# 내부 유틸
# ─────────────────────────────────────────

def _unique_cells(row):
    """병합 셀 제거 후 고유 셀만 반환"""
    seen = set()
    cells = []
    for cell in row.cells:
        if id(cell._tc) not in seen:
            seen.add(id(cell._tc))
            cells.append(cell)
    return cells


def _set_text(cell, value, left_align=False, center=False):
    """텍스트 셀에 값 설정 (기존 run 서식 유지)"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    if value is None:
        value = ''
    value = str(value).strip()
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = value
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(value)

    if left_align:
        # 셀 속성에서 tcFitText / noWrap 제거 (글자 균등배분 원인)
        tcPr = cell._tc.find(qn('w:tcPr'))
        if tcPr is not None:
            for tag in ('w:tcFitText', 'w:noWrap'):
                el = tcPr.find(qn(tag))
                if el is not None:
                    tcPr.remove(el)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_unit(cell, value, unit):
    """단위 셀: 값이 있으면 '167 cm' 형식으로, 없으면 단위만"""
    if value is None or str(value).strip() == '':
        return
    para = cell.paragraphs[0]
    new_text = f'{value} {unit}'
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(new_text)


def _check(cell, options_to_check):
    """
    체크박스 셀에서 선택된 옵션만 ☑, 나머지는 □
    options_to_check: list[str]
    """
    para = cell.paragraphs[0]
    if not para.runs:
        return
    text = para.runs[0].text
    # 기존 ☑ → □ 초기화
    text = text.replace('\u2611', '\u25a1')
    # 선택 옵션 체크
    for opt in options_to_check:
        text = text.replace(f'\u25a1 {opt}', f'\u2611 {opt}')
        text = text.replace(f'\u25a1{opt}', f'\u2611{opt}')
    para.runs[0].text = text


def _replace_colon_fill_balanced_paren(text: str, label: str, content: str) -> str:
    """
    label 직전에 여는 '('부터 짝이 맞는 ')'까지를 한 블록으로 보고,
    그 안에서 label 뒤 첫 ':' 다음~그 블록을 닫는 ')' 직전을 content 한 번으로 바꾼다.
    값에 괄호가 있어도(예: 아스피린) 짝 맞춤으로 닫는 괄호를 찾는다.
    """
    content = (content or '').strip()
    if not content:
        return text
    i = text.find(label)
    if i < 0:
        return text
    colon = text.find(':', i)
    if colon < 0:
        return text
    open_idx = text.rfind('(', 0, i)
    if open_idx < 0:
        return text
    depth = 0
    close_idx = None
    for j in range(open_idx, len(text)):
        c = text[j]
        if c == '(':
            depth += 1
        elif c == ')':
            depth -= 1
            if depth == 0:
                close_idx = j
                break
    if close_idx is None or close_idx <= colon:
        return text
    # ': ' 한 칸 뒤에 내용만 (콜론 뒤·닫는 괄호 앞 공백·옛 값·중복 삽입 제거)
    return text[: colon + 1] + ' ' + content + ' ' + text[close_idx:]


def _check_with_content(cell, selected_option, content=''):
    """
    체크박스 + 내용 텍스트 셀
    예: □ 없음   □ 있음 (내용:                )
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    para = cell.paragraphs[0]
    if not para.runs:
        return
    text = para.runs[0].text
    # 초기화
    text = text.replace('\u2611', '\u25a1')
    # 체크
    if selected_option:
        text = text.replace(f'\u25a1 {selected_option}', f'\u2611 {selected_option}')
        text = text.replace(f'\u25a1{selected_option}', f'\u2611{selected_option}')
    # 내용 삽입 — 라벨별로 콜론 뒤~닫는 괄호 앞을 한 번만 치환 (이전에는 여러 re.sub가 중복 삽입)
    if content:
        content = str(content).strip()
        applied = False
        for label in ('약물명 및 복용 이유', '해당식품', '내용'):
            if label in text:
                new_t = _replace_colon_fill_balanced_paren(text, label, content)
                if new_t != text:
                    text = new_t
                    applied = True
                    break
        if not applied:
            text = re.sub(r'기타\(\s+\)', f'기타({content})', text)
    para.runs[0].text = text
    # 템플릿이 '양쪽 맞춤'·글자 균등 분배면 콜론 뒤 공백이 화면에서 과하게 벌어짐 → 왼쪽 맞춤으로 고정
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ─────────────────────────────────────────
# 메인 채우기 함수
# ─────────────────────────────────────────

_MAX_MEAL_PHOTO_BYTES = 15 * 1024 * 1024
_MEAL_PHOTO_COL = '식사사진등첨부'
_MEAL_PHOTO_COL2 = '식사사진등첨부2'
# 한 장·두 장 모두 세로로 과도하게 늘어나지 않도록 상한 (cm). 2열 나란히 시 각 칸·행 높이 기준.
_MEAL_SINGLE_MAX_W_CM = 11.5
_MEAL_SINGLE_MAX_H_CM = 7.0
_MEAL_PAIR_GAP_CM = 0.45
_MEAL_PAIR_MAX_EACH_W_CM = 6.35
_MEAL_PAIR_MAX_H_CM = 5.8
# ▶ 수급자·보호자 라벨 열: tcW만으로는 tblGrid보다 넓어질 수 없어, 같은 인덱스의 gridCol도 맞춘다.
_INDIVIDUAL_NEEDS_LABEL_GRID_COL = 1
_INDIVIDUAL_NEEDS_LABEL_WIDTH_CM = 1.65


def _cm_to_twips(cm: float) -> int:
    """센티미터 → twips (dxa, Word 1/20 pt)."""
    return int(round(cm * 1440 / 2.54))


def _ensure_table_grid_col_min_twips(doc: Document, grid_index: int, min_twips: int):
    """w:tblGrid의 한 열 너비를 최소 min_twips 이상으로 (실제 열 폭은 gridCol이 지배)."""
    from docx.oxml.ns import qn

    if not doc.tables:
        return
    tbl = doc.tables[0]._tbl
    grid_el = tbl.find(qn('w:tblGrid'))
    if grid_el is None:
        return
    cols = grid_el.findall(qn('w:gridCol'))
    if grid_index < 0 or grid_index >= len(cols):
        return
    min_twips = int(max(100, min_twips))
    c = cols[grid_index]
    w = c.get(qn('w:w'))
    cur = int(w) if w and str(w).isdigit() else 0
    new_w = max(cur, min_twips)
    c.set(qn('w:w'), str(new_w))


def _parse_meal_photo_urls(raw) -> list:
    """시트 셀 값에서 http(s) URL 목록 추출 (=IMAGE / =HYPERLINK 문자열 지원)."""
    if raw is None:
        return []
    s = str(raw).strip()
    if not s:
        return []
    m = re.match(r'^\s*=IMAGE\s*\(\s*["\']([^"\']+)["\']', s, re.I)
    if m:
        s = m.group(1).strip()
    if s.startswith('='):
        hm = re.search(r'HYPERLINK\s*\(\s*"([^"]+)"', s, re.I)
        if hm:
            s = hm.group(1).strip()
        else:
            hm2 = re.search(r"HYPERLINK\s*\(\s*'([^']+)'", s, re.I)
            if hm2:
                s = hm2.group(1).strip()
    out = []
    for part in re.split(r'[\n\r;,]+', s):
        p = part.strip()
        if p.startswith(('http://', 'https://')):
            out.append(p)
    return out


def _cell_clear_to_single_empty_paragraph(cell):
    from docx.oxml.ns import qn

    tc = cell._tc
    for p_el in tc.findall(qn('w:p'))[1:]:
        tc.remove(p_el)
    if not cell.paragraphs:
        cell.add_paragraph()
        return
    p0 = cell.paragraphs[0]
    for r in p0.runs:
        r.text = ''


def _apply_exif_orientation(image_bytes: bytes) -> bytes:
    """
    시트/브라우저는 EXIF Orientation을 반영해 보여주지만, python-docx는 무시해
    가로로 저장된 JPEG이 옆으로 보일 수 있다. 픽셀에 회전을 반영한 뒤 다시 인코딩한다.
    """
    try:
        from PIL import Image, ImageOps
    except ImportError:
        return image_bytes
    try:
        im = Image.open(io.BytesIO(image_bytes))
        im = ImageOps.exif_transpose(im)
        out = io.BytesIO()
        fmt = (im.format or 'JPEG').upper()
        if fmt == 'PNG' or im.mode in ('RGBA', 'LA') or (
            im.mode == 'P' and 'transparency' in im.info
        ):
            im.save(out, format='PNG')
        else:
            im.convert('RGB').save(out, format='JPEG', quality=92, optimize=True)
        return out.getvalue()
    except Exception:
        return image_bytes


def _cell_strip_all_blocks(cell):
    """셀 안의 문단·중첩 표만 제거한다. w:tcPr(gridSpan 등 병합 정보)은 반드시 유지."""
    from docx.oxml.ns import qn

    tc = cell._tc
    block_tags = {qn('w:p'), qn('w:tbl')}
    for child in list(tc):
        if child.tag in block_tags:
            tc.remove(child)


def _meal_image_dims_cm(data: bytes, max_w_cm: float, max_h_cm: float):
    """비율 유지한 채 직사각형 (max_w_cm × max_h_cm) 안에 맞는 가로·세로."""
    from docx.shared import Cm

    try:
        from PIL import Image

        im = Image.open(io.BytesIO(data))
        wpx, hpx = im.size
        if wpx <= 0 or hpx <= 0:
            raise ValueError('bad image size')
    except Exception:
        s = min(max_w_cm, max_h_cm) * 0.55
        return Cm(s), Cm(s)

    aw = max_w_cm
    ah = max_h_cm
    aspect = wpx / hpx
    box = aw / ah
    if aspect > box:
        wc = aw
        hc = aw / aspect
    else:
        hc = ah
        wc = ah * aspect
    return Cm(wc), Cm(hc)


def _insert_meal_photos_cell(cell, raw_primary, raw_secondary=None, config=None):
    """
    식사사진등첨부 칸에 이미지 삽입.
    1열만 있으면 세로로 쌓음. 2열 URL이 모두 있으면 같은 행에 나란히(가운데 간격),
    인덱스별로 짝을 맞춤. 크기는 상한 안에서 자동 축소(2페이지 내 쓰기 목적).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    urls_p = _parse_meal_photo_urls(raw_primary)
    urls_s = (
        _parse_meal_photo_urls(raw_secondary)
        if raw_secondary is not None and str(raw_secondary).strip()
        else []
    )
    if not urls_p and not urls_s:
        return

    import sheets

    try:
        creds = sheets.build_credentials(config)
    except Exception:
        creds = None

    def fetch_one(url):
        try:
            return sheets.fetch_image_bytes(url, creds)
        except Exception:
            try:
                return sheets.fetch_image_bytes(url, None)
            except Exception:
                return None

    imgs_p = []
    for url in urls_p:
        data = fetch_one(url)
        if not data or len(data) > _MAX_MEAL_PHOTO_BYTES:
            continue
        data = _apply_exif_orientation(data)
        imgs_p.append(data)

    imgs_s = []
    for url in urls_s:
        data = fetch_one(url)
        if not data or len(data) > _MAX_MEAL_PHOTO_BYTES:
            continue
        data = _apply_exif_orientation(data)
        imgs_s.append(data)

    if not imgs_p and not imgs_s:
        return

    # 셀 안에 중첩 표를 넣으면 Word가 부모 표의 열 그리드를 다시 잡아 위쪽 행이 깨질 수 있어,
    # 문단 인라인(그림 run 나란히)만 사용한다.
    _cell_strip_all_blocks(cell)

    n = max(len(imgs_p), len(imgs_s))
    has_secondary_column = len(imgs_s) > 0

    for i in range(n):
        left = imgs_p[i] if i < len(imgs_p) else None
        right = imgs_s[i] if i < len(imgs_s) else None

        para = cell.add_paragraph()
        if i > 0:
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(0)

        if left is not None and right is not None and has_secondary_column:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            w1, h1 = _meal_image_dims_cm(left, _MEAL_PAIR_MAX_EACH_W_CM, _MEAL_PAIR_MAX_H_CM)
            w2, h2 = _meal_image_dims_cm(right, _MEAL_PAIR_MAX_EACH_W_CM, _MEAL_PAIR_MAX_H_CM)
            r1 = para.add_run()
            try:
                r1.add_picture(io.BytesIO(left), width=w1, height=h1)
            except Exception:
                pass
            # 고정 폭 공백으로 간격 (중첩 표 없이)
            para.add_run('\u2003\u2003')
            r2 = para.add_run()
            try:
                r2.add_picture(io.BytesIO(right), width=w2, height=h2)
            except Exception:
                pass

        elif left is not None:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            w, h = _meal_image_dims_cm(left, _MEAL_SINGLE_MAX_W_CM, _MEAL_SINGLE_MAX_H_CM)
            run = para.add_run()
            try:
                run.add_picture(io.BytesIO(left), width=w, height=h)
            except Exception:
                pass

        elif right is not None:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            w, h = _meal_image_dims_cm(right, _MEAL_SINGLE_MAX_W_CM, _MEAL_SINGLE_MAX_H_CM)
            run = para.add_run()
            try:
                run.add_picture(io.BytesIO(right), width=w, height=h)
            except Exception:
                pass


def fill_document(data: dict, config=None) -> Document:
    """
    data: 구글 시트 1행 데이터 dict
    config: 선택. 식사사진 등 Drive 다운로드 시 build_credentials에 전달.
    반환: 채워진 Document 객체
    """
    doc = Document(TEMPLATE_PATH)
    table = doc.tables[0]
    rows = {i: _unique_cells(row) for i, row in enumerate(table.rows)}

    def v(key, default=''):
        val = data.get(key, default)
        if val is None:
            return default
        return val

    def is_true(key):
        val = data.get(key, False)
        if isinstance(val, bool):
            return val
        if isinstance(val, str):
            return val.strip().upper() in ('TRUE', '참', '1', 'Y', 'YES')
        return bool(val)

    # ── 행01: 성명 / 입소일 / 생년월일 / 성별 ──
    _set_text(rows[1][1], v('성명'), center=True)
    _set_text(rows[1][3], v('입소일'))
    _set_text(rows[1][5], v('생년월일'))
    _set_text(rows[1][7], v('성별'))

    # ── 행02: 작성일 / 작성자(영양사) ──
    작성일 = v('작성일') or datetime.date.today().strftime('%Y-%m-%d')
    _set_text(rows[2][1], 작성일)
    영양사 = v('영양사이름')
    if 영양사:
        _set_text(rows[2][3], f'{영양사} 영양사')

    # ── 행04: 신장 / 평소체중 / 등급 / 식사유형 ──
    _set_unit(rows[4][1], v('신장'), 'cm')
    _set_unit(rows[4][3], v('평소체중'), 'kg')
    _set_text(rows[4][5], v('등급'))
    _set_text(rows[4][7], v('식사유형'))

    # ── 행06: 1일필요열량 / 1일필요단백질 ──
    _set_unit(rows[6][1], v('1일필요열량'), 'kcal')
    _set_unit(rows[6][3], v('1일필요단백질'), 'g')

    # ── 행08: 식사방법 / 식사섭취상태 (시트: 열별 TRUE 또는 예전 단일 열) ──
    식사방법_목록 = []
    if is_true('식사방법_자립식사'):
        식사방법_목록.append('자립식사')
    if is_true('식사방법_부분도움'):
        식사방법_목록.append('부분도움')
    if is_true('식사방법_완전도움'):
        식사방법_목록.append('완전도움')
    식사방법_레거시 = v('식사방법')
    if 식사방법_목록:
        _check(rows[8][1], 식사방법_목록)
    elif 식사방법_레거시:
        _check(rows[8][1], [식사방법_레거시])

    섭취_목록 = []
    for col, 라벨 in (
        ('식사섭취상태_양호', '양호'),
        ('식사섭취상태_보통', '보통'),
        ('식사섭취상태_불량', '불량'),
    ):
        if is_true(col):
            섭취_목록.append(라벨)
    식사섭취_레거시 = v('식사섭취상태')
    if 섭취_목록:
        _check(rows[8][3], 섭취_목록)
    elif 식사섭취_레거시:
        _check(rows[8][3], [식사섭취_레거시])

    # ── 행09: 식사속도 / 도구사용 (다중) ──
    속도_목록 = []
    for col, 라벨 in (
        ('식사속도_양호', '양호'),
        ('식사속도_보통', '보통'),
        ('식사속도_불량', '불량'),
    ):
        if is_true(col):
            속도_목록.append(라벨)
    식사속도_레거시 = v('식사속도')
    if 속도_목록:
        _check(rows[9][1], 속도_목록)
    elif 식사속도_레거시:
        _check(rows[9][1], [식사속도_레거시])

    도구 = []
    if is_true('도구_젓가락'):     도구.append('젓가락')
    if is_true('도구_숟가락'):     도구.append('숟가락')
    if is_true('도구_포크숟가락'): 도구.append('포크숟가락')
    if is_true('도구_불가'):       도구.append('불가')
    if 도구:
        _check(rows[9][3], 도구)

    # ── 행10: 식사시 문제점 (다중) ──
    문제 = []
    if is_true('문제_식욕저하'): 문제.append('식욕저하')
    if is_true('문제_저작곤란'): 문제.append('저작곤란')
    if is_true('문제_연하곤란'): 문제.append('연하곤란')
    if is_true('문제_소화불량'): 문제.append('소화불량')
    if is_true('문제_구토'):     문제.append('구토')
    if is_true('문제_없음'):     문제.append('없음')
    if 문제:
        _check(rows[10][1], 문제)

    # ── 행12: 치아상태 ──
    치아 = v('치아상태')
    if 치아:
        _check(rows[12][1], [치아])

    # ── 행13: 소화기능 ──
    소화 = v('소화기능')
    if 소화:
        _check(rows[13][1], [소화])

    # ── 행14: 배설양상 ──
    배설 = v('배설양상')
    if 배설:
        _check(rows[14][1], [배설])

    # ── 행15: 특이체질 (시트: 없음/있음 열 + 내용) ──
    특이 = v('특이체질내용')
    if is_true('특이체질_없음'):
        _check(rows[15][1], ['없음'])
    elif is_true('특이체질_있음') or 특이:
        _check_with_content(rows[15][1], '있음', 특이)
    else:
        _check(rows[15][1], ['없음'])

    # ── 행17: 선호음식 / 비선호음식 ──
    _set_text(rows[17][1], v('선호음식', '없음') or '없음')
    _set_text(rows[17][3], v('비선호음식', '없음') or '없음')

    # ── 행18: 식품알러지 ──
    알러지 = v('식품알러지내용')
    if is_true('식품알러지_없음'):
        _check(rows[18][1], ['없음'])
    elif is_true('식품알러지_있음') or 알러지:
        _check_with_content(rows[18][1], '있음', 알러지)
    else:
        _check(rows[18][1], ['없음'])

    # ── 행20: 주요진단명 ──
    _set_text(rows[20][1], v('주요진단명'), left_align=True)

    # ── 행21: 주요질환 (다중) ──
    질환 = []
    if is_true('질환_당뇨'):       질환.append('당뇨')
    if is_true('질환_고혈압'):     질환.append('고혈압')
    if is_true('질환_심장질환'):   질환.append('심장질환')
    if is_true('질환_뇌혈관질환'): 질환.append('뇌혈관질환')
    if is_true('질환_신장질환'):   질환.append('신장질환')
    if is_true('질환_간질환'):     질환.append('간질환')
    if is_true('질환_암'):         질환.append('암')
    기타질환 = v('질환_기타내용')
    if is_true('질환_기타') or 기타질환:
        질환.append('기타')
    if 질환:
        _check(rows[21][1], 질환)
    if 기타질환:
        para = rows[21][1].paragraphs[0]
        if para.runs:
            para.runs[0].text = re.sub(
                r'기타\(\s*\)', f'기타({기타질환})', para.runs[0].text
            )

    # ── 행22: 현재복용약물 (시트: 없음/있음 열 — '악물' 오타 열명도 인식) ──
    약물 = v('복용약물내용')
    복용_없음 = is_true('현재복용약물_없음') or is_true('현재복용악물_없음')
    복용_있음 = is_true('현재복용약물_있음') or is_true('현재복용악물_있음')
    if 복용_있음 and 약물:
        _check_with_content(rows[22][1], '있음', 약물)
    elif 복용_없음:
        _check(rows[22][1], ['없음'])
    elif 약물:
        _check_with_content(rows[22][1], '있음', 약물)
    else:
        _check(rows[22][1], ['없음'])

    # ── 행23: 영양관련약물영향 (다중) ──
    기타약물영향 = v('약물영향_기타내용')
    if is_true('약물영향_없음'):
        약물영향 = ['없음']
    else:
        약물영향 = []
        if is_true('약물영향_식욕저하'):
            약물영향.append('식욕저하')
        if is_true('약물영향_구역구토'):
            약물영향.append('구역/구토')
        if is_true('약물영향_미각변화'):
            약물영향.append('미각변화')
        if is_true('약물영향_흡수장애'):
            약물영향.append('흡수장애')
        if is_true('약물영향_기타') or 기타약물영향:
            약물영향.append('기타')
        if not 약물영향:
            약물영향.append('없음')
    _check(rows[23][1], 약물영향)
    if 기타약물영향:
        para = rows[23][1].paragraphs[0]
        if para.runs:
            para.runs[0].text = re.sub(
                r'기타\(\s*', f'기타({기타약물영향}', para.runs[0].text
            )

    # ── 행25: 종교 / 금식일기도시간 ──
    종교 = v('종교', '없음') or '없음'
    _check(rows[25][1], [종교])
    종교_기타상세 = v('종교_기타내용')
    if 종교_기타상세:
        para = rows[25][1].paragraphs[0]
        if para.runs:
            para.runs[0].text = re.sub(
                r'기타\(\s*\)', f'기타({종교_기타상세})', para.runs[0].text
            )
    _set_text(rows[25][3], v('금식일기도시간'))

    # ── 행26: 종교적식사제한 ──
    종교제한 = v('종교적식사제한', '없음') or '없음'
    종교제한내용 = v('종교제한내용')
    _check_with_content(rows[26][1], 종교제한, 종교제한내용)

    # ── 행27: 문화적식습관 (시트: 플래그 열 + 내용) ──
    문화 = v('문화적식습관내용')
    if is_true('문화적식습관') or 문화:
        _check_with_content(rows[27][1], '있음', 문화)
    else:
        _check(rows[27][1], ['없음'])

    # ── 행28: 출신지역특성 (시트: 출신지역국가특성 + 내용) ──
    출신 = v('출신지역특성내용')
    if is_true('출신지역국가특성') or 출신:
        _check_with_content(rows[28][1], '있음', 출신)
    else:
        _check(rows[28][1], ['해당없음'])

    # ── 행30/31: 개별 욕구 ──
    _set_text(rows[30][2], v('수급자욕구'), left_align=True)
    _set_text(rows[31][2], v('보호자욕구'), left_align=True)

    # ── 행33: 영양사 총평 ──
    _set_text(rows[33][1], v('영양사총평'), left_align=True)

    # ── 행35: 식사사진등첨부 (오른쪽 셀 — 2열이 있으면 같은 줄 나란히, 크기 상한) ──
    photo1 = v(_MEAL_PHOTO_COL, '')
    photo2 = v(_MEAL_PHOTO_COL2, '')
    if str(photo1).strip() or str(photo2).strip():
        _insert_meal_photos_cell(rows[35][1], photo1, photo2, config=config)

    return doc


# ─────────────────────────────────────────
# 저장 / 출력
# ─────────────────────────────────────────

# LibreOffice PDF: 템플릿 라벨에 '성     명', '작 성 일'처럼 넓은 공백이 있으면
# 좁은 셀에서 글자마다 줄바꿈(세로처럼 보임)이 난다. 라벨 열만 공백을 정리한다.
_LO_HANGUL_ADJ_SPACE = re.compile(r'([\uac00-\ud7a3])\s([\uac00-\ud7a3])')


def _compact_label_text_for_libreoffice(s: str) -> str:
    if not s:
        return s
    # 자간용 다중 공백 제거
    t = re.sub(r'\s{2,}', '', s)
    # '작 성 일' → '작성일' (짧은 한글 라벨만, □·숫자·: 등이 있으면 건너뜀)
    stripped = t.strip()
    if len(stripped) > 14:
        return t
    if not re.fullmatch(r'[\uac00-\ud7a3\s]+', stripped):
        return t
    while True:
        u = _LO_HANGUL_ADJ_SPACE.sub(r'\1\2', t)
        if u == t:
            break
        t = u
    return t


def _rewrite_paragraph_text(para, new_text: str):
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ''


def _normalize_lo_label_cells_for_pdf(doc: Document):
    """첫 번째 표의 라벨 열만 LibreOffice 줄바꿈 완화용 텍스트 정리."""
    if not doc.tables:
        return
    table = doc.tables[0]
    for row in table.rows:
        cells = _unique_cells(row)
        n = len(cells)
        if n == 8:
            label_ix = (0, 2, 4, 6)
        elif n == 4:
            label_ix = (0, 2)
        elif n == 2:
            label_ix = (0,)
        elif n == 3:
            label_ix = (0, 1)
        else:
            continue
        for i in label_ix:
            if i >= n:
                continue
            cell = cells[i]
            for para in cell.paragraphs:
                full = ''.join(run.text or '' for run in para.runs)
                new = _compact_label_text_for_libreoffice(full)
                if new != full:
                    _rewrite_paragraph_text(para, new)


def _cell_plain_text(cell) -> str:
    return ''.join(
        (run.text or '')
        for para in cell.paragraphs
        for run in para.runs
    )


def _set_cell_preferred_width_dxa(cell, twips: int):
    """셀 선호 너비(twips, dxa). LibreOffice가 좁은 칸에서 글자 단위 줄바꿈할 때 완화."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    twips = int(max(400, twips))
    tc = cell._tc
    tcPr = tc.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')


def _lo_clear_tc_no_wrap_flags(cell):
    """LibreOffice가 noWrap/tcFitText를 엄격히 적용해 한글 라벨이 세로로 쪼개질 때 제거."""
    from docx.oxml.ns import qn

    tcPr = cell._tc.tcPr
    if tcPr is None:
        return
    for tag in (qn('w:noWrap'), qn('w:tcFitText')):
        el = tcPr.find(tag)
        if el is not None:
            tcPr.remove(el)


def _widen_individual_needs_sublabels_for_lo_pdf(doc: Document):
    """
    ▶ 수급자 및 보호자 개별 욕구: 가운데 '수급자'·'보호자' 라벨 열.
    Word에서는 한 줄로 보이나 LibreOffice PDF는 열을 더 좁게 잡아 글자 단위 줄바꿈이 나는 경우가 있다.
    tcW(셀 선호 너비)만 올리면 tblGrid의 해당 세로줄이 더 좁으면 효과가 없으므로,
    gridCol과 tcW를 같은 twips(1.65cm)로 맞춘다.
    """
    if not doc.tables:
        return
    table = doc.tables[0]
    label_cells = []
    for row in table.rows:
        cells = _unique_cells(row)
        if len(cells) != 3:
            continue
        raw = _cell_plain_text(cells[1])
        compact = re.sub(r'\s+', '', (raw or '').strip())
        if compact not in ('수급자', '보호자'):
            continue
        label_cells.append(cells[1])
    if not label_cells:
        return
    tw = _cm_to_twips(_INDIVIDUAL_NEEDS_LABEL_WIDTH_CM)
    _ensure_table_grid_col_min_twips(doc, _INDIVIDUAL_NEEDS_LABEL_GRID_COL, tw)
    for cell in label_cells:
        _lo_clear_tc_no_wrap_flags(cell)
        _set_cell_preferred_width_dxa(cell, tw)


def _twips_for_label_cell(text: str) -> int:
    """라벨 글자 수에 따른 목표 너비(twips). LibreOffice 첫 열용—과하면 값 칸이 좁아진다."""
    t = (text or '').strip()
    n = max(2, len(t))
    # 이전(2800~·320/글자) 대비 약 절반: 짧은 라벨은 최소만, 긴 라벨만 비례 확대
    return min(2400, max(1380, 155 * n + 380))


def _widen_first_column_labels_for_lo_pdf(doc: Document):
    """
    각 데이터 행의 첫 번째 셀(라벨 열)만 너비 확대.
    PDF에서만 세로로 깨지는 경우가 많아 Linux 저장 시에만 호출한다.
    """
    from docx.oxml.ns import qn

    if not doc.tables:
        return
    table = doc.tables[0]
    for row in table.rows:
        cells = _unique_cells(row)
        n = len(cells)
        if n < 2 or n not in (2, 3, 4, 8):
            continue
        cell0 = cells[0]
        raw = _cell_plain_text(cell0)
        if not raw.strip():
            continue
        compact = _compact_label_text_for_libreoffice(raw)
        w = _twips_for_label_cell(compact)
        cur = cell0._tc.tcPr
        if cur is not None:
            tcw = cur.find(qn('w:tcW'))
            if tcw is not None:
                cur_w = tcw.get(qn('w:w'))
                if cur_w and str(cur_w).isdigit() and int(cur_w) >= w:
                    continue
        _set_cell_preferred_width_dxa(cell0, w)


def _rebalance_tbl_grid_first_col_for_lo_pdf(doc: Document):
    """
    LibreOffice는 w:tblGrid의 첫 w:gridCol(템플릿 약 461 twips)을 엄격히 적용해
    첫 라벨이 글자 단위로 세로 배치된다. Word 로컬에서는 덜 드러날 수 있음.

    첫 gridCol 목표 폭을 라벨 길이에 맞추고, 다른 열에서는 **최소 폭을 지키며**
    줄일 수 있는 만큼만 빼서 총폭을 유지한다.
    """
    from docx.oxml.ns import qn

    if not doc.tables:
        return
    tbl = doc.tables[0]._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        return
    cols = grid.findall(qn('w:gridCol'))
    if len(cols) < 2:
        return

    widths = []
    for c in cols:
        w = c.get(qn('w:w'))
        widths.append(int(w) if w and str(w).isdigit() else 0)

    # 첫 열 라벨(그리드 1칸만 쓰는 행) 중 필요한 최대 폭 (_twips_for_label_cell과 동일 하한)
    target0 = 1380
    table = doc.tables[0]
    for row in table.rows:
        cells = _unique_cells(row)
        if len(cells) < 2 or len(cells) not in (2, 3, 4, 8):
            continue
        tc0 = cells[0]._tc
        tcPr = tc0.tcPr
        if tcPr is None:
            continue
        gs = tcPr.find(qn('w:gridSpan'))
        if gs is not None and gs.get(qn('w:val')) != '1':
            continue
        raw = _cell_plain_text(cells[0])
        if not raw.strip():
            continue
        compact = _compact_label_text_for_libreoffice(raw)
        target0 = max(target0, _twips_for_label_cell(compact))

    old0 = widths[0]
    # 다른 열에서 각각 최소 이 정도는 남긴다 (생년월일·등급 값 칸이 세로로 가는 것 방지)
    min_floor = 620
    max_stealable = sum(max(0, w - min_floor) for w in widths[1:])
    want = target0 - old0
    if want <= 0:
        return
    delta = min(want, max_stealable)
    if delta <= 0:
        return
    target0 = old0 + delta
    remaining = delta
    order = sorted(range(1, len(widths)), key=lambda i: widths[i], reverse=True)
    while remaining > 0:
        progressed = False
        for i in order:
            if remaining <= 0:
                break
            take = min(remaining, max(0, widths[i] - min_floor))
            if take > 0:
                widths[i] -= take
                remaining -= take
                progressed = True
        if not progressed:
            break
    widths[0] = old0 + (delta - remaining)

    for c, w in zip(cols, widths):
        c.set(qn('w:w'), str(int(max(100, w))))


def save_document(doc: Document, name: str) -> str:
    """output/ 폴더에 저장 후 경로 반환"""
    if platform.system() != 'Windows':
        # 라벨 문단 공백만 정리 (LibreOffice PDF 줄바꿈 완화). tblGrid/tcW는 건드리지 않음 —
        # 예전에는 첫 열·gridCol을 다시 잡는 함수가 있었는데, 템플릿에서 맞춘 열 너비(성명 행 등)가
        # Streamlit(Linux) 저장 시에만 덮어씌워져 로컬(Windows)과 달라지는 문제가 있었다.
        _normalize_lo_label_cells_for_pdf(doc)
        _widen_individual_needs_sublabels_for_lo_pdf(doc)
        if os.environ.get('SOON_DOCX_LO_GRID_FIX', '').strip() in ('1', 'true', 'yes'):
            _widen_first_column_labels_for_lo_pdf(doc)
            _rebalance_tbl_grid_first_col_for_lo_pdf(doc)
    # 파일 이름에 사용 불가 문자 제거
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', name)
    path = os.path.join(OUTPUT_DIR, f'{safe_name}.docx')
    doc.save(path)
    return path


def open_for_preview(path: str):
    """첫 번째 문서: Word로 열어서 미리보기"""
    import subprocess
    subprocess.Popen(['cmd', '/c', 'start', '', path], shell=False)


def print_document(path: str):
    """자동 프린트 (Windows 기본 프린터)"""
    import subprocess
    subprocess.Popen(
        ['cmd', '/c', 'start', '/MIN', '', '/PRINT', path],
        shell=False
    )


def generate_all(records: list, config=None) -> list:
    """
    모든 입소자 docx 생성 후 경로 목록 반환
    records: get_all_records() 결과
    """
    paths = []
    for rec in records:
        name = rec.get('성명', '미입력')
        doc = fill_document(rec, config=config)
        path = save_document(doc, name)
        paths.append(path)
    return paths


# ─────────────────────────────────────────
# ZIP + 합본 PDF 생성
# ─────────────────────────────────────────

def build_zip(paths: list) -> bytes:
    """개인별 docx 파일들을 ZIP으로 묶어 bytes 반환"""
    import io
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, os.path.basename(path))
    return buf.getvalue()


def _build_pdf_windows(paths: list, tmp_dir: str) -> list:
    """Windows: Word COM(docx2pdf)으로 PDF 변환, PDF 경로 목록 반환"""
    import pythoncom
    from docx2pdf import convert

    pythoncom.CoInitialize()
    try:
        pdf_paths = []
        for docx_path in paths:
            pdf_path = os.path.join(
                tmp_dir,
                os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            )
            convert(docx_path, pdf_path)
            pdf_paths.append(pdf_path)
        return pdf_paths
    finally:
        pythoncom.CoUninitialize()


def _build_pdf_libreoffice(paths: list, tmp_dir: str) -> list:
    """Linux/클라우드: LibreOffice headless로 PDF 변환, PDF 경로 목록 반환"""
    import subprocess

    pdf_paths = []
    for docx_path in paths:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', tmp_dir, docx_path],
            check=True,
        )
        pdf_path = os.path.join(
            tmp_dir,
            os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        )
        pdf_paths.append(pdf_path)
    return pdf_paths


def build_merged_pdf(paths: list) -> bytes:
    """
    개인별 docx → PDF 변환 후 하나로 합쳐 bytes 반환
    Windows: Word COM(docx2pdf) 사용
    Linux/클라우드: LibreOffice headless 사용
    """
    import io
    from pypdf import PdfWriter, PdfReader

    tmp_dir = tempfile.mkdtemp()
    try:
        if platform.system() == 'Windows':
            pdf_paths = _build_pdf_windows(paths, tmp_dir)
        else:
            pdf_paths = _build_pdf_libreoffice(paths, tmp_dir)

        writer = PdfWriter()
        for pdf_path in pdf_paths:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                writer.add_page(page)

        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue()
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
